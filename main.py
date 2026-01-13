from dotenv import load_dotenv
from google import genai
from google.genai import types
import pathlib
import docx
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

load_dotenv()

print("loading gemini")
client = genai.Client()

# get meetings
print("getting meetings")
meetings_image = pathlib.Path('meetings.png')
with open('meetings_prompt.txt', 'r') as file:
    prompt = file.read()

meetings = client.models.generate_content(
  model="gemini-2.5-flash",
  contents=[
      types.Part.from_bytes(
        data=meetings_image.read_bytes(),
        mime_type='image/png',
      ),
      prompt]).text

# get outings
print("getting outings")
outings_image = pathlib.Path('outings.png')
with open('outings_prompt.txt', 'r') as file:
    prompt = file.read()

outings = client.models.generate_content(
  model="gemini-2.5-flash",
  contents=[
      types.Part.from_bytes(
        data=outings_image.read_bytes(),
        mime_type='image/png',
      ),
      prompt]).text

print(outings)
print(meetings)

# build docx
doc = Document()

# styling
style_normal = doc.styles['Normal']
style_heading2 = doc.styles['Heading 2']
style_heading1 = doc.styles['Heading 1']
style_normal.font.name = 'Arial'
style_heading2.font.name = 'Arial'
style_heading1.font.name = 'Arial'
style_heading2.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
style_heading1.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)
style_normal.paragraph_format.line_spacing = 1
style_heading2.font.size = Pt(16)
style_heading1.font.size = Pt(20)
style_heading2.font.bold = False
style_heading1.font.bold = False

doc.add_paragraph(f'PLC {datetime.today().strftime("%-m/%-d/%Y")}\nPresent:\nAbsent:\n')

# outings
outings_split = outings.strip().split('\n')
for line in outings_split:
    # strip
    line = line.strip()
    if not line:
        continue
    doc.add_paragraph(line, style='Heading 2')
    doc.add_paragraph("", style='Normal')

# meetings
doc.add_paragraph("Meeting Planning", style='Heading 1')
meetings_split = meetings.strip().split('\n')
for line in meetings_split:
    # strip
    line = line.strip()
    if not line:
        continue
    
    # bullet point
    if line.split()[0] in ("Opening", "Skill", "Game", "Intrapatrol", "Closing"):
        doc.add_paragraph(line, style='List Bullet')
    else:
        doc.add_paragraph(line)

doc.save("plc_notes.docx")
print("Outputted to plc_notes.docx")